"""
CAM Generator — generates Credit Appraisal Memo in DOCX and PDF formats.

Produces a comprehensive memo covering all required sections:
1. Executive Summary
2. Borrower Profile
3. Financial Analysis
4. Fraud Detection Findings
5. Promoter Risk Analysis
6. Sector Outlook
7. Five Cs Credit Assessment
8. Loan Recommendation
9. Risk Justification

Exports as DOCX (python-docx) and PDF (reportlab).
"""

import json
import logging
from datetime import datetime
from pathlib import Path

from langchain_core.prompts import ChatPromptTemplate

from app.core.llm import get_recommendation_llm
from app.config import CAM_OUTPUT_DIR
from app.schemas.recommendation import (
    CAMRequest,
    CAMSection,
    CreditAppraisalMemo,
)

logger = logging.getLogger(__name__)

CAM_PROMPT = ChatPromptTemplate.from_messages([
    (
        "system",
        "You are a senior credit officer at an Indian bank. Write a professional, "
        "concise narrative for a Credit Appraisal Memo section. Use formal banking "
        "language. Be specific with numbers and findings. "
        "Indian regulatory context (RBI, SEBI, MCA) is important.",
    ),
    (
        "human",
        """Write the "{section_title}" section for the Credit Appraisal Memo of {company_name}.

Available data:
{section_data}

Write 3-5 concise paragraphs. Be factual and analytical. Include specific numbers where available.
""",
    ),
])

# All required CAM sections
CAM_SECTIONS = [
    "Executive Summary",
    "Borrower Profile",
    "Financial Analysis",
    "Working Capital Stress Analysis",
    "Fraud Detection Summary",
    "Promoter Network Analysis",
    "Sector Risk Intelligence",
    "Early Warning Signals",
    "Borrower Historical Analysis",
    "Five Cs Credit Assessment",
    "Explainable ML Decision",
    "Loan Recommendation",
    "Risk Justification",
]


def generate_cam(request: CAMRequest) -> CreditAppraisalMemo:
    """
    Generate a full Credit Appraisal Memo with DOCX and PDF export.
    """
    sections: list[CAMSection] = []

    # Map sections to their data sources
    section_data_map = {
        "Executive Summary": {
            "loan_decision": request.loan_decision,
            "five_cs_scores": request.five_cs_scores,
            "financial_data": request.financial_data,
        },
        "Borrower Profile": {
            "financial_data": request.financial_data,
            "research_report": request.research_report,
        },
        "Financial Analysis": {
            "financial_data": request.financial_data,
            "cross_verification": request.cross_verification,
        },
        "Working Capital Stress Analysis": {
            "working_capital": request.working_capital,
            "financial_data": request.financial_data,
        },
        "Fraud Detection Summary": {
            "cross_verification": request.cross_verification,
        },
        "Promoter Network Analysis": {
            "promoter_network": request.promoter_network,
            "research_report": request.research_report,
        },
        "Sector Risk Intelligence": {
            "sector_risk": request.sector_risk,
            "research_report": request.research_report,
        },
        "Early Warning Signals": {
            "early_warning": request.early_warning,
        },
        "Borrower Historical Analysis": {
            "historical_trust": request.historical_trust,
        },
        "Five Cs Credit Assessment": {
            "five_cs_scores": request.five_cs_scores,
        },
        "Explainable ML Decision": {
            "explainability": request.loan_decision.get("explainability", {}),
            "ml_risk_prediction": request.loan_decision.get("ml_risk_prediction", {}),
            "final_credit_risk_score": request.loan_decision.get("final_credit_risk_score", 0.0),
        },
        "Loan Recommendation": {
            "loan_decision": request.loan_decision,
            "five_cs_scores": request.five_cs_scores,
        },
        "Risk Justification": {
            "loan_decision": request.loan_decision,
            "explainability": request.loan_decision.get("explainability", {}),
        },
    }

    for title in CAM_SECTIONS:
        data = section_data_map.get(title, {})
        content = _generate_section(request.company_name, title, data)
        sections.append(CAMSection(title=title, content=content))

    # Extract recommendation details
    decision = request.loan_decision
    recommendation = decision.get("decision", "REFER")
    recommended_amount = decision.get("recommended_amount", 0.0)
    interest_rate = decision.get("interest_rate", 0.0)
    risk_grade = decision.get("risk_grade", "")

    # Generate DOCX and PDF
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_name = request.company_name.replace(" ", "_")[:30]

    docx_path = _generate_docx(
        request.company_name, sections, decision, timestamp, safe_name,
    )
    pdf_path = _generate_pdf(
        request.company_name, sections, decision, timestamp, safe_name,
    )

    return CreditAppraisalMemo(
        company_name=request.company_name,
        generated_at=datetime.now().isoformat(),
        sections=sections,
        executive_summary=sections[0].content if sections else "",
        recommendation=recommendation,
        risk_grade=risk_grade,
        recommended_amount=recommended_amount,
        interest_rate=interest_rate,
        docx_path=docx_path,
        pdf_path=pdf_path,
    )


def _generate_section(company_name: str, title: str, data: dict) -> str:
    """Generate a single CAM section using LLM."""
    try:
        llm = get_recommendation_llm()
        chain = CAM_PROMPT | llm

        section_data = json.dumps(data, indent=2, default=str)[:3000]

        result = chain.invoke({
            "company_name": company_name,
            "section_title": title,
            "section_data": section_data,
        })

        return result.content if hasattr(result, "content") else str(result)

    except Exception as e:
        logger.error("CAM section '%s' generation failed: %s", title, e)
        # Generate a data-based fallback
        return _fallback_section(title, data)


def _fallback_section(title: str, data: dict) -> str:
    """Generate a basic section from data when LLM fails."""
    parts = [f"## {title}\n"]

    if not data:
        parts.append("No data available for this section.")
        return "\n".join(parts)

    for key, value in data.items():
        if isinstance(value, dict):
            parts.append(f"\n### {key.replace('_', ' ').title()}")
            for k, v in value.items():
                if v is not None:
                    parts.append(f"- {k.replace('_', ' ').title()}: {v}")
        elif isinstance(value, list):
            parts.append(f"\n### {key.replace('_', ' ').title()}")
            for item in value[:5]:
                parts.append(f"- {item}")
        elif value is not None:
            parts.append(f"- {key.replace('_', ' ').title()}: {value}")

    return "\n".join(parts)


def _generate_docx(
    company_name: str,
    sections: list[CAMSection],
    decision: dict,
    timestamp: str,
    safe_name: str,
) -> str:
    """Generate DOCX file using python-docx."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Title
        title = doc.add_heading(
            "CREDIT APPRAISAL MEMO", level=0,
        )
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Company name
        company_para = doc.add_heading(company_name, level=1)
        company_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Date
        date_para = doc.add_paragraph(
            f"Generated: {datetime.now().strftime('%d %B %Y, %H:%M')}",
        )
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Decision summary box
        doc.add_heading("Decision Summary", level=2)
        table = doc.add_table(rows=4, cols=2)
        table.style = "Light Grid Accent 1"
        cells = [
            ("Decision", decision.get("decision", "REFER")),
            ("Risk Grade", decision.get("risk_grade", "N/A")),
            ("Recommended Amount", f"₹{decision.get('recommended_amount', 0):,.0f}"),
            ("Interest Rate", f"{decision.get('interest_rate', 0):.2f}%"),
        ]
        for i, (label, value) in enumerate(cells):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = str(value)

        doc.add_paragraph("")

        # Sections
        for section in sections:
            doc.add_heading(section.title, level=2)
            for paragraph in section.content.split("\n"):
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())

        # Save
        CAM_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
        filename = f"CAM_{safe_name}_{timestamp}.docx"
        filepath = CAM_OUTPUT_DIR / filename
        doc.save(str(filepath))

        logger.info("DOCX generated: %s", filepath)
        return str(filepath)

    except ImportError:
        logger.warning("python-docx not available — DOCX generation skipped")
        return ""
    except Exception as e:
        logger.error("DOCX generation failed: %s", e)
        return ""


def _generate_pdf(
    company_name: str,
    sections: list[CAMSection],
    decision: dict,
    timestamp: str,
    safe_name: str,
) -> str:
    """Generate PDF file using reportlab."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib.colors import HexColor
        from reportlab.platypus import (
            SimpleDocTemplate,
            Paragraph,
            Spacer,
            Table,
            TableStyle,
        )

        CAM_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
        filename = f"CAM_{safe_name}_{timestamp}.pdf"
        filepath = CAM_OUTPUT_DIR / filename

        doc = SimpleDocTemplate(
            str(filepath),
            pagesize=A4,
            topMargin=0.75 * inch,
            bottomMargin=0.75 * inch,
        )

        styles = getSampleStyleSheet()

        # Custom styles
        title_style = ParagraphStyle(
            "CAMTitle",
            parent=styles["Title"],
            fontSize=20,
            spaceAfter=6,
            textColor=HexColor("#1a237e"),
        )
        heading_style = ParagraphStyle(
            "CAMHeading",
            parent=styles["Heading2"],
            fontSize=14,
            spaceAfter=6,
            spaceBefore=12,
            textColor=HexColor("#283593"),
        )
        body_style = ParagraphStyle(
            "CAMBody",
            parent=styles["Normal"],
            fontSize=10,
            spaceAfter=6,
            leading=14,
        )

        elements = []

        # Title
        elements.append(Paragraph("CREDIT APPRAISAL MEMO", title_style))
        elements.append(Paragraph(company_name, heading_style))
        elements.append(Paragraph(
            f"Generated: {datetime.now().strftime('%d %B %Y, %H:%M')}",
            body_style,
        ))
        elements.append(Spacer(1, 12))

        # Decision table
        elements.append(Paragraph("Decision Summary", heading_style))
        table_data = [
            ["Parameter", "Value"],
            ["Decision", decision.get("decision", "REFER")],
            ["Risk Grade", decision.get("risk_grade", "N/A")],
            ["Recommended Amount", f"₹{decision.get('recommended_amount', 0):,.0f}"],
            ["Interest Rate", f"{decision.get('interest_rate', 0):.2f}%"],
        ]
        t = Table(table_data, colWidths=[200, 250])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), HexColor("#1a237e")),
            ("TEXTCOLOR", (0, 0), (-1, 0), HexColor("#ffffff")),
            ("ALIGN", (0, 0), (-1, -1), "LEFT"),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 10),
            ("BOTTOMPADDING", (0, 0), (-1, 0), 8),
            ("TOPPADDING", (0, 0), (-1, -1), 6),
            ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#cccccc")),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [
                HexColor("#f5f5f5"),
                HexColor("#ffffff"),
            ]),
        ]))
        elements.append(t)
        elements.append(Spacer(1, 12))

        # Sections
        for section in sections:
            elements.append(Paragraph(section.title, heading_style))
            for para in section.content.split("\n"):
                if para.strip():
                    # Escape XML special chars
                    safe_para = (
                        para.strip()
                        .replace("&", "&amp;")
                        .replace("<", "&lt;")
                        .replace(">", "&gt;")
                    )
                    elements.append(Paragraph(safe_para, body_style))
            elements.append(Spacer(1, 8))

        doc.build(elements)
        logger.info("PDF generated: %s", filepath)
        return str(filepath)

    except ImportError:
        logger.warning("reportlab not available — PDF generation skipped")
        return ""
    except Exception as e:
        logger.error("PDF generation failed: %s", e)
        return ""
